"""Generate WIP Document docx for Matundu SMS"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(os.path.abspath(__file__))
STUDENT = "Mwange Musa Muthami"
REG = "23/05037"
SUPERVISOR = "Dr. Simon N. Mwendia / Dr. Kevin Mugoye Sindu"
TITLE = "School Management System for Matundu Primary School"
DATE = "August 2026"

def new_doc():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.5
    for sec in d.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(2.54)
    return d

def cover(d, doc_title):
    for _ in range(6): d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE.upper()); r.bold = True; r.font.size = Pt(16)
    d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_title.upper()); r.bold = True; r.font.size = Pt(14)
    for text in [STUDENT, f"Registration Number: {REG}", f"Submission Date: {DATE}", f"Supervisor: {SUPERVISOR}",
                 "A project document submitted in partial fulfillment of the requirements for Bachelor of Information Technology (BIT 3105)"]:
        d.add_paragraph()
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = "submitted" not in text
    d.add_page_break()

def h1(d, t): d.add_heading(t, level=1)
def h2(d, t): d.add_heading(t, level=2)
def para(d, t): d.add_paragraph(t)
def bullet(d, t): d.add_paragraph(t, style='List Bullet')

def add_table(d, headers, rows):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row): rc[i].text = str(v)
    d.add_paragraph()

d = new_doc()
cover(d, "Work in Progress Document")

h1(d, "DECLARATION")
para(d, "I declare that this work in progress document is my original work and has not been submitted elsewhere for examination.")
para(d, "Signature: ___________________________       Date: ________________________")
para(d, f"{STUDENT} ({REG})")
para(d, "\nSupervisor Approval:")
para(d, "Signature: ___________________________       Date: ________________________")
para(d, SUPERVISOR)
d.add_page_break()

h1(d, "TABLE OF CONTENTS")
para(d, "(Auto-generate from Word: References > Table of Contents)")
d.add_page_break()

h1(d, "1. PROJECT OVERVIEW")
h2(d, "1.1 Project Title")
para(d, "School Management System for Matundu Primary School")
h2(d, "1.2 Student Details")
add_table(d, ["Field", "Detail"],
    [["Student Name", STUDENT], ["Registration Number", REG],
     ["Course", "Bachelor of Information Technology (BIT 3105)"],
     ["Supervisors", SUPERVISOR], ["Academic Year", "2025/2026"],
     ["Institution", "KCA University"]])
h2(d, "1.3 Project Description")
para(d, "The Matundu Primary School Management System (MATUNDU-SMS) is a production-grade, full-stack web application designed to digitize and streamline all administrative and academic operations at Matundu Primary School in Voo, Kitui County. The system is fully aligned with Kenya's Competency-Based Curriculum (CBC) and handles the complete lifecycle of school operations including student admissions, attendance tracking, CBC-aligned assessments, report card generation, fee management, timetabling, and stakeholder communication.")
h2(d, "1.4 Problem Being Solved")
para(d, "Matundu Primary School has historically relied on paper-based record keeping for all administrative functions. This creates critical challenges:")
for p in ["Data loss and corruption of physical student records",
          "Manual computation of CBC competency grades consuming excessive teacher time",
          "No real-time access for parents to track student progress and fee balances",
          "Slow and error-prone fee reconciliation by the bursar",
          "Inability to generate standardized CBC-compliant report cards quickly"]:
    bullet(d, p)
d.add_page_break()

h1(d, "2. TECHNOLOGY STACK")
add_table(d, ["Layer", "Technology", "Version"],
    [["Frontend", "React 18 + Vite + Tailwind CSS", "18.3 / 7.3 / 3.4"],
     ["State Management", "Zustand", "5.0.1"],
     ["Charts", "Recharts", "2.13.3"],
     ["HTTP Client", "Axios", "1.7.7"],
     ["Router", "React Router DOM", "6.28.0"],
     ["PDF (Frontend)", "jsPDF + AutoTable", "4.1.0 / 5.0.7"],
     ["Backend", "Node.js 20 + Express.js", "20+ / 4.21.1"],
     ["ORM", "Prisma", "5.22.0"],
     ["Database", "PostgreSQL", "15+"],
     ["Cache", "Redis", "5.10.0"],
     ["Auth", "JWT + bcrypt", "—"],
     ["PDF (Backend)", "PDFKit", "0.15.0"],
     ["Logging", "Winston", "3.15.0"],
     ["Security", "Helmet", "8.0.0"],
     ["Containerization", "Docker + Compose", "—"],
     ["Deployment (FE)", "Cloudflare Pages", "—"],
     ["Deployment (BE)", "Render", "—"]])
d.add_page_break()

h1(d, "3. PROJECT ARCHITECTURE")
h2(d, "3.1 System Architecture Overview")
para(d, "The system employs a three-tier client-server architecture deployed on modern cloud infrastructure:")
add_table(d, ["Tier", "Technology", "Hosting", "Purpose"],
    [["Client (Frontend)", "React 18 SPA", "Cloudflare Pages", "User interface for all roles"],
     ["Server (Backend)", "Express.js REST API", "Render", "Business logic and data access"],
     ["Database", "PostgreSQL 15", "Render (managed)", "Persistent relational data storage"],
     ["Cache", "Redis", "Render (managed)", "Rate limiting and session caching"]])
h2(d, "3.2 Project Folder Structure")
para(d, "matundu-sms/")
for line in ["  ├── backend/src/app.js             # Express entry point (32KB)",
             "  ├── backend/src/config/             # DB, Redis, env config",
             "  ├── backend/src/middleware/         # Auth, RBAC, audit, sanitize",
             "  ├── backend/src/features/           # 13 feature modules",
             "  ├── backend/prisma/schema.prisma    # 24-model DB schema",
             "  ├── frontend/src/App.jsx            # React router root",
             "  ├── frontend/src/pages/             # 17 page components",
             "  ├── frontend/src/components/        # Reusable components",
             "  ├── frontend/src/stores/            # Zustand global state",
             "  ├── frontend/src/services/          # Axios API wrappers",
             "  ├── docker-compose.yml              # Local dev environment",
             "  └── render.yaml                     # Cloud deployment config"]:
    para(d, line)
d.add_page_break()

h1(d, "4. DATABASE DESIGN")
h2(d, "4.1 Schema Overview")
para(d, "The Prisma schema defines 24 models and 6 enums covering all functional domains of the school management system.")
h2(d, "4.2 Core Models")
add_table(d, ["Model", "Purpose", "Key Fields"],
    [["User", "Base authentication entity", "id, email, password, role, isActive"],
     ["Student", "Learner profiles", "admissionNumber, firstName, classId, admissionStatus"],
     ["Staff", "Teacher/admin profiles", "employeeNumber, firstName, qualification, specialization"],
     ["Guardian", "Parent/guardian details", "firstName, lastName, relationship, nationalId"],
     ["AcademicYear", "School year definition", "name, startDate, endDate, isCurrent"],
     ["Term", "Term within academic year", "termNumber, startDate, endDate, academicYearId"],
     ["Grade", "CBC grades PP1-Grade 9", "name, level"],
     ["Stream", "Class streams", "name"],
     ["Class", "Grade+Stream+Year combo", "name, gradeId, streamId, academicYearId"],
     ["Subject", "Subject definitions", "name, code, gradeId"],
     ["Attendance", "Daily attendance records", "date, status, studentId, classId, termId"],
     ["Assessment", "Test/exam definitions", "name, type, maxScore, weight, subjectId, termId"],
     ["AssessmentScore", "Student scores", "score, grade, studentId, assessmentId"],
     ["Competency", "CBC competencies", "name, description, subjectId"],
     ["CompetencyScore", "CBC ratings per student", "rating (EE/ME/AE/BE), studentId, competencyId"],
     ["ReportCard", "Termly report", "totalScore, averageScore, rank, pdfUrl"],
     ["FeeStructure", "Fee config per grade/term", "name, amount, gradeId, termId"],
     ["StudentInvoice", "Student billing", "invoiceNo, totalAmount, paidAmount, balance"],
     ["Payment", "Payment records", "amount, method, status, transactionRef, mpesaReceiptNo"],
     ["Announcement", "School announcements", "title, content, targetRoles, isPublished"],
     ["Message", "Direct messages", "subject, content, senderId, receiverId"],
     ["TimetableSlot", "Weekly timetable", "dayOfWeek, startTime, endTime, classId, subjectId"],
     ["AuditLog", "System audit trail", "action, entity, userId, ipAddress"],
     ["AiChatMessage", "AI assistant history", "role, content, studentId"]])
h2(d, "4.3 Enumerations")
add_table(d, ["Enum", "Values"],
    [["Role", "SUPER_ADMIN, ADMIN, TEACHER, STUDENT, PARENT, BURSAR"],
     ["AdmissionStatus", "PENDING, APPROVED, REJECTED"],
     ["AttendanceStatus", "PRESENT, ABSENT, LATE, EXCUSED"],
     ["CompetencyRating", "EXCEEDING (EE), MEETING (ME), APPROACHING (AE), BELOW (BE)"],
     ["PaymentMethod", "CASH, MPESA, BANK_TRANSFER"],
     ["PaymentStatus", "PENDING, COMPLETED, FAILED, REFUNDED"]])
d.add_page_break()

h1(d, "5. BACKEND IMPLEMENTATION")
h2(d, "5.1 Middleware Pipeline")
add_table(d, ["Order", "Middleware", "File", "Purpose"],
    [["1", "CORS", "cors", "Cross-origin request handling"],
     ["2", "Helmet", "helmet.js", "Security headers (XSS, CSRF, clickjacking)"],
     ["3", "Rate Limiter", "express-rate-limit + Redis", "Brute-force protection"],
     ["4", "Request ID", "requestId.js", "Unique request tracking"],
     ["5", "Body Parser", "express.json()", "JSON request parsing"],
     ["6", "Sanitize", "sanitize.js", "Strip HTML/script injection"],
     ["7", "Auth (JWT)", "auth.js", "JWT token verification"],
     ["8", "RBAC", "rbac.js", "Role-based permission enforcement"],
     ["9", "Validate", "validate.js", "express-validator result handling"],
     ["10", "Audit Log", "auditLog.js", "Mutation activity recording"],
     ["11", "Error Handler", "errorHandler.js", "Global error formatting"]])
h2(d, "5.2 Feature Modules Implemented")
add_table(d, ["Module", "Files", "Key Functionality"],
    [["auth", "4 files", "Login, register, JWT refresh, profile, password change"],
     ["students", "4 files", "CRUD, admission workflow, guardian linking, class assignment"],
     ["staff", "3 files", "Staff CRUD, teacher-class-subject assignments"],
     ["academic", "12 files", "Years, terms, grades, classes, subjects, timetable, outlines, resources"],
     ["assessments", "3 files", "Create assessments, bulk score entry, CBC grading"],
     ["attendance", "5 files", "Daily attendance marking, QR code attendance service"],
     ["reports", "3 files", "Report card generation, class rankings, PDF export"],
     ["fees", "4 files", "Fee structures, invoice generation, payments, M-Pesa service"],
     ["communication", "3 files", "Announcements with role targeting, direct messaging"],
     ["ai", "3 files", "AI chat assistant service with conversation history"],
     ["reminders", "3 files", "Student reminders with Google/Microsoft calendar sync"],
     ["integrations", "3 files", "Google Calendar OAuth, Microsoft Tasks via Graph API"],
     ["dashboard", "—", "Aggregated KPIs for admin and student dashboards"]])
d.add_page_break()

h1(d, "6. FRONTEND IMPLEMENTATION")
h2(d, "6.1 Pages Implemented")
add_table(d, ["Page", "File Size", "Role Access", "Description"],
    [["Login", "3.9 KB", "Public", "JWT authentication form"],
     ["Dashboard", "40.2 KB", "All Roles", "Analytics hub with KPI cards and Recharts"],
     ["Students", "11.9 KB", "Admin/Teacher", "Student list with search and filter"],
     ["StudentDetail", "20.3 KB", "Admin/Teacher", "Full student profile, attendance, scores, fees"],
     ["Staff", "11.8 KB", "Admin", "Staff management and assignment"],
     ["Classes", "14.7 KB", "Admin/Teacher", "Class and academic structure management"],
     ["Assessments", "21.6 KB", "Teacher", "Score entry with auto-grading"],
     ["Attendance", "10.8 KB", "Teacher", "Daily attendance marking interface"],
     ["Fees", "21.3 KB", "Bursar/Admin", "Invoice management and payment recording"],
     ["Reports", "50.4 KB", "Admin/Teacher", "Report card generation with PDF export"],
     ["Timetable", "22.1 KB", "Admin/Teacher", "Visual weekly timetable grid"],
     ["Announcements", "12.9 KB", "All", "School-wide announcements"],
     ["Admissions", "10.2 KB", "Admin", "Admission approval/rejection workflow"],
     ["Settings", "22.7 KB", "Admin", "System configuration"],
     ["Profile", "3.7 KB", "All", "User profile management"],
     ["SystemStatus", "6.7 KB", "Admin", "Backend health monitoring"]])
d.add_page_break()

h1(d, "7. FEATURE COMPLETION STATUS")
add_table(d, ["Feature Module", "Backend", "Frontend", "Status", "Notes"],
    [["Authentication & RBAC", "✅ Done", "✅ Done", "Complete", "JWT + 6 roles enforced"],
     ["Student Management", "✅ Done", "✅ Done", "Complete", "Full CRUD + admissions"],
     ["Staff Management", "✅ Done", "✅ Done", "Complete", "—"],
     ["Academic Structure (CBC)", "✅ Done", "✅ Done", "Complete", "PP1–Grade 9"],
     ["Timetable", "✅ Done", "✅ Done", "Complete", "Visual weekly grid"],
     ["Course Outlines", "✅ Done", "Partial", "In Progress", "Backend done; UI in academic pages"],
     ["Course Resources", "✅ Done", "Partial", "In Progress", "File upload backend ready"],
     ["Attendance (Daily)", "✅ Done", "✅ Done", "Complete", "4 statuses"],
     ["Attendance (QR Code)", "✅ Done", "Partial", "In Progress", "Frontend UI incomplete"],
     ["Assessments & Grading", "✅ Done", "✅ Done", "Complete", "CBC + traditional grades"],
     ["Report Cards", "✅ Done", "✅ Done", "Complete", "PDF export available"],
     ["Fee Management", "✅ Done", "✅ Done", "Complete", "Invoicing + payments"],
     ["M-Pesa STK Push", "✅ Done", "Partial", "In Progress", "Backend service ready; not tested E2E"],
     ["Announcements", "✅ Done", "✅ Done", "Complete", "Role-targeted"],
     ["Direct Messaging", "✅ Done", "Partial", "In Progress", "Basic UI"],
     ["AI Chat Assistant", "✅ Done", "Partial", "In Progress", "Needs API key config"],
     ["Reminders", "✅ Done", "Partial", "In Progress", "Linked to Google/MS"],
     ["Google Calendar", "✅ Done", "Partial", "In Progress", "OAuth wired"],
     ["Microsoft Tasks", "✅ Done", "Partial", "In Progress", "MS Graph API wired"],
     ["Audit Logging", "✅ Done", "N/A", "Complete", "Backend-only"],
     ["Dashboard Analytics", "✅ Done", "✅ Done", "Complete", "KPIs, charts"],
     ["Flutter Parent App", "Not Started", "Not Started", "Pending", "Future milestone"],
     ["SMS Notifications", "Not Started", "Not Started", "Pending", "Future milestone"],
     ["Automated Test Suite", "Not Started", "Not Started", "Pending", "Vitest configured"]])
d.add_page_break()

h1(d, "8. PROGRESS SUMMARY BY PHASE")
h2(d, "Phase 1 — Foundation (COMPLETE ✅)")
for item in ["Database schema designed (24 models, 6 enums)", "Express app configured with full middleware stack",
             "JWT authentication with refresh token rotation", "Role-based access control on all 6 roles",
             "Student management module (CRUD + admissions workflow)", "Staff management module", "Academic structure (grades, classes, subjects)"]:
    bullet(d, item)
h2(d, "Phase 2 — Core Academic (COMPLETE ✅)")
for item in ["Timetable management with weekly grid view", "Attendance tracking (daily marking)", "Assessment and CBC grading engine",
             "CBC competency scoring (EE/ME/AE/BE)", "Report card generation with rankings", "Course outlines (JSON-based weekly plan)", "Course resources (file upload)"]:
    bullet(d, item)
h2(d, "Phase 3 — Finance & Communication (COMPLETE ✅)")
for item in ["Fee structures per grade per term", "Invoice generation with unique numbering", "Payment recording (Cash, M-Pesa, Bank Transfer)",
             "M-Pesa STK Push service (backend ready)", "Announcements with role targeting", "Direct messaging between users"]:
    bullet(d, item)
h2(d, "Phase 4 — Advanced Features (IN PROGRESS 🟡)")
for item in ["AI chat assistant backend ✅ | Frontend UI partial",
             "QR code attendance backend ✅ | Frontend UI incomplete",
             "Google Calendar OAuth integration ✅ | Frontend partial",
             "Microsoft Tasks integration ✅ | Frontend partial",
             "Cloud file storage for resources ❌ (needs Cloudflare R2 or S3)"]:
    bullet(d, item)
h2(d, "Phase 5 — Testing & Deployment (IN PROGRESS 🟡)")
for item in ["Docker Compose for local development ✅", "Render deployment config ✅", "Cloudflare Pages config ✅",
             "Automated test suite ❌ (Vitest configured, no tests written)", "CI/CD pipeline ❌ (GitHub Actions not set up)",
             "Production smoke testing ❌ (pending)"]:
    bullet(d, item)
h2(d, "Phase 6 — Mobile & Expansion (NOT STARTED ❌)")
for item in ["Flutter parent mobile application", "SMS/Email notification service",
             "Multi-school tenancy support", "Advanced analytics and reporting"]:
    bullet(d, item)
d.add_page_break()

h1(d, "9. API ENDPOINTS SUMMARY")
h2(d, "Authentication — /api/auth")
add_table(d, ["Method", "Endpoint", "Description"],
    [["POST", "/register", "Register new user account"],
     ["POST", "/login", "Login with email/password, receive JWT tokens"],
     ["POST", "/refresh", "Refresh access token using refresh token"],
     ["GET", "/profile", "Get authenticated user's profile"],
     ["PATCH", "/profile", "Update profile information"],
     ["POST", "/change-password", "Update account password"]])
h2(d, "Students — /api/students")
add_table(d, ["Method", "Endpoint", "Description"],
    [["GET", "/", "List all students with pagination and filters"],
     ["POST", "/", "Create student profile and user account"],
     ["GET", "/:id", "Get full student details"],
     ["PUT", "/:id", "Update student information"],
     ["DELETE", "/:id", "Remove student record"],
     ["POST", "/:id/approve", "Approve student admission"],
     ["POST", "/:id/reject", "Reject student admission"],
     ["POST", "/:id/promote", "Promote student to next grade"]])
h2(d, "Academic — /api/academic")
add_table(d, ["Method", "Endpoint", "Description"],
    [["GET", "/years", "List academic years"],
     ["POST", "/years", "Create academic year"],
     ["GET", "/terms", "List terms"],
     ["GET", "/grades", "List CBC grades (PP1–Grade 9)"],
     ["GET", "/classes", "List classes"],
     ["POST", "/classes", "Create a class"],
     ["GET", "/subjects", "List subjects"],
     ["POST", "/timetable", "Create timetable slot"],
     ["GET", "/timetable", "Get timetable slots"]])
h2(d, "Fees — /api/fees")
add_table(d, ["Method", "Endpoint", "Description"],
    [["GET", "/structures", "List fee structures"],
     ["POST", "/structures", "Create fee structure"],
     ["POST", "/invoices", "Generate student invoice"],
     ["GET", "/invoices/:studentId", "Get student invoices"],
     ["POST", "/payments", "Record a payment"],
     ["POST", "/mpesa/initiate", "Initiate M-Pesa STK Push"],
     ["POST", "/mpesa/callback", "Handle M-Pesa payment callback"]])
d.add_page_break()

h1(d, "10. DEPLOYMENT CONFIGURATION")
h2(d, "10.1 Backend — Render.com")
add_table(d, ["Setting", "Value"],
    [["Service Name", "matundu-sms-backend"],
     ["Runtime", "Node.js 20"],
     ["Region", "Oregon (us-west)"],
     ["Build Command", "npm install && npx prisma generate"],
     ["Start Command", "npx prisma migrate deploy && npm start"],
     ["Database", "PostgreSQL (matundu-db, free tier)"],
     ["Cache", "Redis (matundu-redis, free tier)"],
     ["CORS Origin", "https://matundu-frontend.pages.dev"],
     ["Auth Rate Limit", "15 requests max"]])
h2(d, "10.2 Frontend — Cloudflare Pages")
add_table(d, ["Setting", "Value"],
    [["Config File", "frontend/wrangler.toml"],
     ["Build Command", "npm run build"],
     ["Output Directory", "dist"],
     ["Target URL", "https://matundu-frontend.pages.dev"]])
h2(d, "10.3 Local Development — Docker Compose")
para(d, "Run 'docker-compose up -d' to start all services:")
add_table(d, ["Service", "Port", "Description"],
    [["PostgreSQL", "5432", "Database server"],
     ["Backend API", "3000", "Express.js API server"],
     ["Frontend", "5173", "Vite dev server"]])
d.add_page_break()

h1(d, "11. CHALLENGES AND SOLUTIONS")
add_table(d, ["#", "Challenge Encountered", "Solution Applied"],
    [["1", "Supabase connection failures in production", "Configured Prisma connection pooling and retry logic"],
     ["2", "Auto-seed running on every login causing data loss", "Moved seeding to one-time setup; removed from login flow"],
     ["3", "Assessment scores not grouped by term in frontend", "Refactored API query to include term relations; grouped in UI"],
     ["4", "Dashboard showing hardcoded placeholder values", "Replaced all static values with live API calls"],
     ["5", "Redis rate limiting not working in free-tier Render", "Added graceful fallback to memory-based rate limiting"],
     ["6", "QR code attendance UI incomplete", "Backend service complete; frontend UI deferred to Phase 4"],
     ["7", "M-Pesa callback URL not accessible in local dev", "Using ngrok for local callback testing; prod URL configured"],
     ["8", "Google/Microsoft OAuth redirect URIs mismatch", "Documented required URIs for production deployment config"]])
d.add_page_break()

h1(d, "12. KNOWN ISSUES AND BLOCKERS")
add_table(d, ["Priority", "Issue", "Area", "Status"],
    [["HIGH", "M-Pesa STK Push not tested end-to-end in production", "Fees", "Pending"],
     ["HIGH", "AI chat service requires external API key configuration", "AI Module", "Pending"],
     ["HIGH", "No automated test suite written (Vitest configured only)", "Testing", "Not Started"],
     ["HIGH", "No CI/CD pipeline configured in GitHub Actions", "DevOps", "Not Started"],
     ["MEDIUM", "Google/Microsoft OAuth redirect URIs need production update", "Integrations", "Pending"],
     ["MEDIUM", "QR Code attendance frontend UI is incomplete", "Attendance", "In Progress"],
     ["MEDIUM", "Course Resources need cloud storage (Cloudflare R2 or S3)", "Academic", "Pending"],
     ["MEDIUM", "Redis rate limiting needs validation in production environment", "Security", "Pending"],
     ["LOW", "Timetable.jsx_snippet file is a leftover debug artifact", "Frontend", "Cleanup Needed"],
     ["LOW", "PDF report card layout needs design polish", "Reports", "Pending"]])
d.add_page_break()

h1(d, "13. NEXT STEPS")
h2(d, "Immediate (Week 1–2)")
for item in ["Write automated unit tests using Vitest for all service modules",
             "Complete QR code attendance frontend scanning UI",
             "Set up GitHub Actions CI/CD workflow for auto-deploy on push",
             "Configure production OAuth redirect URIs for Google and Microsoft",
             "Polish AI chat student-facing interface"]:
    bullet(d, item)
h2(d, "Short Term (Week 3–4)")
for item in ["End-to-end M-Pesa STK Push sandbox testing",
             "Integrate Cloudflare R2 for course resource file storage",
             "Complete Swagger/OpenAPI documentation for all endpoints",
             "Full production smoke test on Render + Cloudflare Pages"]:
    bullet(d, item)
h2(d, "Long Term")
for item in ["Develop Flutter parent mobile application",
             "Integrate Africa's Talking or similar SMS gateway for notifications",
             "Database query optimization and API performance profiling",
             "Security audit and penetration testing",
             "Extend architecture for multi-school tenancy support"]:
    bullet(d, item)
d.add_page_break()

h1(d, "14. REFERENCES")
for ref in ["Ministry of Education. (2023). National Guidelines on Basic Education Information Management Systems. Republic of Kenya.",
            "Pressman, R. S. (2014). Software Engineering: A Practitioner's Approach (8th ed.). McGraw-Hill.",
            "Prisma. (2024). Prisma ORM Documentation. Retrieved from https://www.prisma.io/docs",
            "React. (2024). React 18 Documentation. Retrieved from https://react.dev",
            "Render. (2024). Deploy Node.js Applications. Retrieved from https://render.com/docs",
            "ShulePro. (2021). School Administration Software Overview. Retrieved from https://www.shulepro.com",
            "Sommerville, I. (2015). Software Engineering (10th ed.). Pearson.",
            "Zeraki Analytics. (2023). Transforming Education through Data. Retrieved from https://zeraki.co.ke"]:
    para(d, ref)

path = os.path.join(OUT, "9_Work_In_Progress.docx")
d.save(path)
print(f"✓ Saved: {path}")
